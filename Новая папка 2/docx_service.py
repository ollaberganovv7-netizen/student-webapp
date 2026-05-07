import io, re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SERVICE_TITLES = {
    "essay": "REFERAT",
    "report": "MUSTAQIL ISH",
    "coursework": "KURS ISHI",
    "tezis": "ILMIY MAQOLA (TEZIS)",
    "maqola": "ILMIY MAQOLA",
}

def _set_heading_style(paragraph, level: int = 1):
    """Apply standard academic heading formatting."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        run.font.size = Pt(16)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        run.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(12)

def generate_docx(service_type: str, topic: str, content: str, author: str = "Talaba", plan: str = "", extra_meta: dict = None) -> bytes:
    doc = Document()
    extra_meta = extra_meta or {}

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)

    # ── Default style ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.first_line_indent = Cm(1.25)

    # ── COVER PAGE ───────────────────────────────────────────────────────────
    if service_type == "coursework":
        # Professional Coursework Cover
        min_p = doc.add_paragraph(extra_meta.get("ministry", "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM VAZIRLIGI").upper())
        min_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        min_p.paragraph_format.first_line_indent = Cm(0)
        min_p.runs[0].font.size = Pt(12)
        min_p.runs[0].font.bold = True

        for _ in range(3): doc.add_paragraph("")

        kaf_p = doc.add_paragraph(extra_meta.get("department", "KAFEDRA NOMI").upper())
        kaf_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kaf_p.paragraph_format.first_line_indent = Cm(0)
        kaf_p.runs[0].font.size = Pt(14)
        kaf_p.runs[0].font.bold = True

        for _ in range(4): doc.add_paragraph("")

        p_type = doc.add_paragraph("KURS ISHI")
        p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_type.paragraph_format.first_line_indent = Cm(0)
        p_type.runs[0].font.size = Pt(28)
        p_type.runs[0].font.bold = True

        p_top = doc.add_paragraph(f"MAVZU: {topic.upper()}")
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.first_line_indent = Cm(0)
        p_top.runs[0].font.size = Pt(16)
        p_top.runs[0].font.italic = True

        for _ in range(6): doc.add_paragraph("")

        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_auth.paragraph_format.first_line_indent = Cm(0)
        run_auth = p_auth.add_run(f"Bajardi: {author}\nTekshirdi: ________________")
        run_auth.font.size = Pt(14)
        run_auth.font.bold = True

    elif service_type == "tezis":
        # Tezis title page: topic bold caps centered, author right italic
        p_top = doc.add_paragraph(topic.upper())
        p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_top.paragraph_format.first_line_indent = Cm(0)
        p_top.runs[0].font.size = Pt(16)
        p_top.runs[0].font.bold = True

        p_auth = doc.add_paragraph(f"Muallif: {author}")
        p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_auth.paragraph_format.first_line_indent = Cm(0)
        p_auth.runs[0].font.size = Pt(12)
        p_auth.runs[0].font.italic = True

        doc.add_paragraph("")  # Space before text

    elif service_type == "maqola":
        # Maqola title page: matches academic journal format
        # 1) Title — bold, caps, centered
        p_title = doc.add_paragraph(topic.upper())
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.first_line_indent = Cm(0)
        p_title.paragraph_format.space_after = Pt(6)
        run_title = p_title.runs[0]
        run_title.font.name = "Times New Roman"
        run_title.font.size = Pt(14)
        run_title.font.bold = True

        # 2) Author — bold, centered
        p_auth = doc.add_paragraph(author)
        p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_auth.paragraph_format.first_line_indent = Cm(0)
        p_auth.paragraph_format.space_after = Pt(2)
        run_auth = p_auth.runs[0]
        run_auth.font.name = "Times New Roman"
        run_auth.font.size = Pt(14)
        run_auth.font.bold = True

        # 3) University — italic, centered (only if provided)
        university = extra_meta.get("university", "")
        if university:
            p_uni = doc.add_paragraph(university)
            p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_uni.paragraph_format.first_line_indent = Cm(0)
            p_uni.paragraph_format.space_after = Pt(2)
            run_uni = p_uni.runs[0]
            run_uni.font.name = "Times New Roman"
            run_uni.font.size = Pt(14)
            run_uni.font.italic = True

        doc.add_paragraph("")  # Space before content

    elif service_type == "uslubiy":
        # Load from uslubiy template if it exists
        import os
        template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "doc_templates", "uslubiy_template.docx")
        if os.path.exists(template_path):
            doc = Document(template_path)
            university = extra_meta.get("university", "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM VAZIRLIGI").upper()
            import datetime
            year = str(datetime.datetime.now().year)

            for p in doc.paragraphs:
                text = p.text
                if "Universitet" in text and "___" in text:
                    for run in p.runs:
                        if "___" in run.text:
                            run.text = run.text.replace("_______________________________________________", university)
                            break
                elif "FAN" in text and "___" in text:
                    for run in p.runs:
                        if "___" in run.text:
                            subject = extra_meta.get("subject", "")
                            if subject:
                                run.text = run.text.replace("_________________", subject.upper())
                            break
                elif "1-MAVZU:" in text and "___" in text:
                    for run in p.runs:
                        if "___" in run.text:
                            run.text = run.text.replace("____________________________________", topic.upper())
                            break
                elif "202" in text and "yil" in text:
                    for run in p.runs:
                        if "202" in run.text:
                            run.text = run.text.replace("2026", year)
                            break
        else:
            # Fallback if template is missing
            doc.add_paragraph("USLUBIY ISHLANMA").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"MAVZU: {topic.upper()}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    else:
        # Standard Cover (Referat/Mustaqil ish)
        doc_type_label = SERVICE_TITLES.get(service_type, "REFERAT")
        p1 = doc.add_paragraph("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM VAZIRLIGI")
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.first_line_indent = Cm(0)
        p1.runs[0].font.size = Pt(12)
        p1.runs[0].font.bold = True
        for _ in range(5): doc.add_paragraph("")
        p2 = doc.add_paragraph(doc_type_label)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.first_line_indent = Cm(0)
        p2.runs[0].font.size = Pt(28)
        p2.runs[0].font.bold = True
        p3 = doc.add_paragraph(f"Mavzu: {topic}")
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.first_line_indent = Cm(0)
        p3.runs[0].font.size = Pt(16)
        p3.runs[0].font.italic = True
        for _ in range(8): doc.add_paragraph("")
        p_auth = doc.add_paragraph()
        p_auth.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_auth.paragraph_format.first_line_indent = Cm(0)
        run_auth = p_auth.add_run(f"Bajardi: {author}")
        run_auth.font.size = Pt(14)
        run_auth.font.bold = True

    if service_type != "tezis":
        doc.add_page_break()

    # ── MUNDARIJA ──────────────────────────────────────────────────────────
    if plan:
        p_mun = doc.add_paragraph("MUNDARIJA")
        _set_heading_style(p_mun, level=1)
        p_mun.paragraph_format.first_line_indent = Cm(0)
        plan_lines = plan.split("\n")
        for line in plan_lines:
            clean_line = line.strip().lstrip("0123456789. ")
            if clean_line:
                p = doc.add_paragraph(clean_line)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.2
                p.runs[0].font.size = Pt(14)
        doc.add_page_break()

    # ── CONTENT ─────────────────────────────────────────────────────────────
    lines = content.split("\n")
    i_line = 0
    while i_line < len(lines):
        stripped = lines[i_line].strip()
        if not stripped:
            i_line += 1
            continue

        if stripped.startswith("#"):
            level = 1 if stripped.startswith("# ") else 2
            p = doc.add_paragraph(stripped.lstrip("#").strip())
            _set_heading_style(p, level=level)
            p.paragraph_format.first_line_indent = Cm(0)
            i_line += 1
            continue

        # ── Image placeholder: [ 🖼️ SHU YERGA RASM JOYLANG: ... ]
        if "🖼" in stripped and "SHU YERGA RASM JOYLANG" in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(stripped)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i_line += 1
            continue

        # ── Table: JADVAL: header followed by | delimited rows
        if stripped.upper().startswith("JADVAL:") or ("|" in stripped and stripped.count("|") >= 2):
            # Collect all table rows
            table_rows = []
            if stripped.upper().startswith("JADVAL:"):
                i_line += 1  # skip "JADVAL:" label
            while i_line < len(lines):
                row_text = lines[i_line].strip()
                if not row_text or ("|" not in row_text and row_text.count("|") < 1):
                    break
                # Skip separator lines like |---|---|
                if re.match(r'^[\s|:-]+$', row_text):
                    i_line += 1
                    continue
                cells = [c.strip() for c in row_text.split("|") if c.strip() != ""]
                if cells:
                    table_rows.append(cells)
                i_line += 1

            if table_rows:
                max_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.style = 'Table Grid'
                for r_idx, row_data in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row_data):
                        if c_idx < max_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = cell_text
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.first_line_indent = Cm(0)
                                for run in paragraph.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(12)
                                    if r_idx == 0:
                                        run.bold = True
                doc.add_paragraph("")  # Space after table
            continue

        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        i_line += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Template-based DOCX generation (Referat) ────────────────────────────────

import os

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "doc_templates")

def generate_docx_from_template(
    topic: str,
    content: str,
    author: str = "Talaba",
    plan: str = "",
    subject: str = "",
    reviewer: str = "",
) -> bytes:
    """
    Generate a referat DOCX using the ready-made template file.
    Opens the template, replaces title-page placeholders, fills plan names,
    then removes all empty paragraphs between content headers and inserts
    AI-generated text as new paragraphs with proper formatting.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_BREAK
    from copy import deepcopy

    template_path = os.path.join(TEMPLATE_DIR, "referat_template.docx")
    doc = Document(template_path)

    # ── Parse AI content into ordered list ──────────────────────────────
    sections_list = []  # [(name, text), ...]
    current_section = None
    current_lines = []

    for line in content.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            if current_section is not None:
                sections_list.append((current_section, "\n".join(current_lines)))
            current_section = stripped[2:].strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_section is not None:
        sections_list.append((current_section, "\n".join(current_lines)))

    # ── Parse plan into clean section names ─────────────────────────────
    plan_sections = [s.strip() for s in plan.split("\n") if s.strip()] if plan else []
    # Clean AI numbering from plan sections ("1. Topic name" -> "Topic name")
    clean_plan_sections = []
    for s in plan_sections:
        cleaned = re.sub(r'^[\d\.]+\s*', '', s).strip()
        clean_plan_sections.append(cleaned)
        
    # Extract only sub-section names (not Kirish, Xulosa, Adabiyotlar, Asosiy)
    clean_sub_sections = [s for s in clean_plan_sections if not any(
        s.upper().startswith(k) for k in ["KIRISH", "XULOSA", "FOYDALANILGAN", "ASOSIY"]
    )]

    slot_keys = ["1.1.", "1.2.", "2.1.", "2.2."]
    plan_replacements = {}
    for i, key in enumerate(slot_keys):
        if i < len(clean_sub_sections):
            plan_replacements[key] = f"{key} {clean_sub_sections[i]}"

    # ── Replace title page placeholders ─────────────────────────────────
    for p in doc.paragraphs:
        text = p.text
        # Fan nomi
        if "FANIDAN" in text and "___" in text:
            for run in p.runs:
                if "___" in run.text:
                    run.text = run.text.replace("_________________________", subject.upper() if subject else "________________")
                    break

        # Mavzu
        if "Mavzu:" in text and "___" in text:
            for run in p.runs:
                if "___" in run.text:
                    run.text = topic
                    break

        # Topshirdi
        if "Topshirdi:" in text and "___" in text:
            for run in p.runs:
                if "___" in run.text:
                    run.text = run.text.replace("_____________________", author)
                    break

        # Qabul qildi
        if "Qabul qildi:" in text and "___" in text:
            for run in p.runs:
                if "___" in run.text:
                    run.text = run.text.replace("____________________", reviewer if reviewer else "________________")
                    break

    # ── Replace plan page section names (paras 17-25 area) ──────────────
    for p in doc.paragraphs[:30]:  # Only first 30 paragraphs (plan page)
        text = p.text.strip()
        for key, new_text in plan_replacements.items():
            if text.startswith(key) and "___" in text:
                for run in p.runs:
                    if key in run.text or "___" in run.text:
                        run.text = new_text
                break

    # ── Identify content section headers (dynamically find start) ───────
    paragraphs = list(doc.paragraphs)

    # Find first KIRISH header — everything before it is title+plan pages
    content_start = 0
    for i, p in enumerate(paragraphs):
        if p.text.strip() == "KIRISH" and i > 25:
            content_start = i
            break

    section_markers = []  # [(para_index, marker_key), ...]
    for i, p in enumerate(paragraphs):
        if i < content_start:
            continue
        text = p.text.strip()
        if not text:
            continue
        if text == "KIRISH":
            section_markers.append((i, "KIRISH"))
        elif text == "ASOSIY QISM":
            section_markers.append((i, "ASOSIY_QISM"))
        elif text.startswith("1.1."):
            section_markers.append((i, "1.1"))
        elif text.startswith("1.2."):
            section_markers.append((i, "1.2"))
        elif text.startswith("2.1."):
            section_markers.append((i, "2.1"))
        elif text.startswith("2.2."):
            section_markers.append((i, "2.2"))
        elif text == "XULOSA":
            section_markers.append((i, "XULOSA"))
        elif "adabiyotlar" in text.lower() or "Foydalanilgan" in text:
            section_markers.append((i, "ADABIYOTLAR"))

    # ── Delete all empty paragraphs between content sections ────────────
    if section_markers:
        first_content_para = section_markers[0][0]
        # Collect all empty paragraph elements to delete (from content area only)
        to_delete = []
        for i in range(first_content_para, len(paragraphs)):
            p = paragraphs[i]
            if not p.text.strip():
                to_delete.append(p._element)

        for elem in to_delete:
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

    # ── Update section header text with plan names ──────────────────────
    # Re-read paragraphs after deletion (empty paras gone, headers are now adjacent)
    paragraphs = list(doc.paragraphs)
    # Re-find section markers — scan all, but only match known content headers
    section_markers = []
    found_kirish = False
    for i, p in enumerate(paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # Once we find the first KIRISH, we're in the content area
        if text == "KIRISH" and not found_kirish and i > 20:
            found_kirish = True
            section_markers.append((i, "KIRISH"))
        elif not found_kirish:
            continue
        elif text == "ASOSIY QISM":
            p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
            section_markers.append((i, "ASOSIY_QISM"))
        elif text.startswith("1.1."):
            section_markers.append((i, "1.1"))
        elif text.startswith("1.2."):
            section_markers.append((i, "1.2"))
        elif text.startswith("2.1."):
            section_markers.append((i, "2.1"))
        elif text.startswith("2.2."):
            section_markers.append((i, "2.2"))
        elif text == "XULOSA":
            p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
            section_markers.append((i, "XULOSA"))
        elif "adabiyotlar" in text.lower() or "Foydalanilgan" in text:
            p.insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
            section_markers.append((i, "ADABIYOTLAR"))

    # Update numbered section headers with plan names
    for para_idx, marker_key in section_markers:
        if marker_key in ("1.1", "1.2", "2.1", "2.2"):
            slot_key = marker_key + "."
            if slot_key in plan_replacements:
                p = paragraphs[para_idx]
                for run in p.runs:
                    if "___" in run.text or slot_key in run.text:
                        run.text = plan_replacements[slot_key]
                        break

    # ── Map sections smartly ────────────────────────────────────────────
    content_markers = [(idx, key) for idx, key in section_markers if key != "ASOSIY_QISM"]
    
    mapped_ai_content = {}
    remaining_ai_sections = []
    
    for name, text in sections_list:
        name_upper = name.upper()
        if "KIRISH" in name_upper:
            mapped_ai_content["KIRISH"] = text
        elif "XULOSA" in name_upper:
            mapped_ai_content["XULOSA"] = text
        elif "ADABIYOTLAR" in name_upper or "FOYDALANILGAN" in name_upper:
            mapped_ai_content["ADABIYOTLAR"] = text
        else:
            remaining_ai_sections.append(text)
            
    # Assign remaining to numbered slots
    numbered_slots = ["1.1", "1.2", "2.1", "2.2"]
    for i, slot in enumerate(numbered_slots):
        if i < len(remaining_ai_sections):
            mapped_ai_content[slot] = remaining_ai_sections[i]

    # ── Insert AI text after each section header ────────────────────────
    # We work backwards to avoid index shifting
    insertions = []  # [(para_element, lines_to_insert), ...]

    for para_idx, marker_key in content_markers:
        ai_text = mapped_ai_content.get(marker_key, "")
        if not ai_text:
            continue

        text_lines = [l.strip() for l in ai_text.strip().split("\n") if l.strip()]
        # Filter out markdown headers
        text_lines = [l for l in text_lines if not l.startswith("#")]

        # Filter out lines that duplicate the section header name
        header_text = paragraphs[para_idx].text.strip().upper()
        filtered_lines = []
        for l in text_lines:
            l_upper = l.upper().strip()
            # Skip if line is just the section name (KIRISH, XULOSA, etc.)
            if l_upper == header_text:
                continue
            if l_upper in ("KIRISH", "XULOSA", "ASOSIY QISM", "FOYDALANILGAN ADABIYOTLAR", "FOYDALANILGAN ADABIYOTLAR RO'YXATI"):
                continue
            # Skip lines like "XULOSA VA TAVSIYALAR" that duplicate section intent
            if marker_key == "XULOSA" and l_upper.startswith("XULOSA"):
                continue
            filtered_lines.append(l)
        text_lines = filtered_lines

        if text_lines:
            header_element = paragraphs[para_idx]._element
            insertions.append((header_element, text_lines))

    # Insert in reverse order to preserve element positions
    for header_element, text_lines in reversed(insertions):
        parent = header_element.getparent()
        insert_after = header_element

        for line_text in text_lines:
            # Create new paragraph element
            new_p = OxmlElement('w:p')

            # Copy paragraph properties from template normal style
            pPr = OxmlElement('w:pPr')
            # Set justified alignment
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'both')
            pPr.append(jc)
            # Set first line indent (1.25cm = 709 twips)
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLine'), '709')
            pPr.append(ind)
            new_p.append(pPr)

            # Check if this is an image placeholder
            is_image_placeholder = "🖼" in line_text and "SHU YERGA RASM JOYLANG" in line_text

            if is_image_placeholder:
                # Center alignment for image placeholders
                jc.set(qn('w:val'), 'center')
                # Remove first line indent
                ind.set(qn('w:firstLine'), '0')

                run_elem = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                # Font
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rPr.append(rFonts)
                # Size 12pt (24 half-points)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '24')
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '24')
                rPr.append(szCs)
                # Bold
                b = OxmlElement('w:b')
                rPr.append(b)
                # Italic
                italic = OxmlElement('w:i')
                rPr.append(italic)
                # Gray color
                color = OxmlElement('w:color')
                color.set(qn('w:val'), '808080')
                rPr.append(color)

                run_elem.append(rPr)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = line_text
                run_elem.append(t)
                new_p.append(run_elem)
            else:
                # Normal text - handle **bold** markers
                clean_parts = re.split(r"(\*\*[^*]+\*\*)", line_text)
                for part in clean_parts:
                    if not part:
                        continue
                    run_elem = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    # Font
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), 'Times New Roman')
                    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                    rPr.append(rFonts)
                    # Size 14pt (28 half-points)
                    sz_elem = OxmlElement('w:sz')
                    sz_elem.set(qn('w:val'), '28')
                    rPr.append(sz_elem)
                    szCs = OxmlElement('w:szCs')
                    szCs.set(qn('w:val'), '28')
                    rPr.append(szCs)

                    if part.startswith("**") and part.endswith("**"):
                        b = OxmlElement('w:b')
                        rPr.append(b)
                        actual_text = part[2:-2]
                    else:
                        actual_text = part

                    run_elem.append(rPr)
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')
                    t.text = actual_text
                    run_elem.append(t)
                    new_p.append(run_elem)

            # Insert after the current position
            insert_after.addnext(new_p)
            insert_after = new_p

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
